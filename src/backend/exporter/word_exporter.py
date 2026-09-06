#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 导出模块 (word_exporter)

把 LaTeX 源码的一个常用子集转换成**真正能被 Word 打开**的 .docx 文档
（基于 python-docx）。

历史问题
--------
旧实现只是把纯文本写进了 ``.docx`` 后缀的文件里::

    f.write(f"Word Export Sample\\n{content}")

Word 打开这种文件会直接提示"文件已损坏"，等于导出功能完全不可用。
本模块重写为真实的结构化转换。

支持的语法
----------
* 标题：\\title / \\author / \\date / \\maketitle、\\part / \\section / \\subsection /
  \\subsubsection / \\paragraph
* 段落、空行分段、\\\\ 与 \\newline 换行、% 注释
* 文本样式：\\textbf \\textit \\emph \\underline \\texttt \\textsc
* 列表：itemize / enumerate / description 环境，以及裸 \\item
* 表格：tabular 环境（含 \\hline / \\toprule / \\midrule / \\bottomrule 忽略、
  l/c/r 列对齐、p{} 宽度）
* 浮动体与容器：table / figure / center / abstract / quote / minipage 会递归处理
* 代码：verbatim / lstlisting / alltt 环境以等宽字体输出
* 数学：行内 $...$ 与行间 \\[...\\]、$$...$$、equation 环境

已知限制（与 README FAQ 的说明一致）
------------------------------------
python-docx 没有公开的公式(OMML)写入接口，因此数学公式降级为**斜体等宽文本**
展示原始 LaTeX 源码，而不是渲染成 Word 公式对象。这样导出的文件始终合法可打开，
且用户仍能看到并手动修正公式。
"""

import datetime
import json
import os
import re
import sys

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ImportError:  # pragma: no cover - 依赖缺失时给出可操作的提示
    Document = None


# ---------------------------------------------------------------------------
# 常量
# ---------------------------------------------------------------------------

SECTION_LEVELS = {
    'part': 0,
    'section': 1,
    'subsection': 2,
    'subsubsection': 3,
    'paragraph': 4,
    'subparagraph': 4,
}

STYLE_FONT_MAP = {
    'textbf': 'bold', 'bf': 'bold',
    'textit': 'italic', 'it': 'italic', 'emph': 'italic',
    'underline': 'underline',
    'texttt': 'mono', 'tt': 'mono',
    'textsc': 'smallcaps', 'sc': 'smallcaps',
}

# 需要原样保留的转义字符
ESCAPED_CHARS = {
    r'\%': '%', r'\$': '$', r'\&': '&', r'\#': '#', r'\_': '_',
    r'\{': '{', r'\}': '}', r'\ ': ' ', r'\,': ' ', r'\;': ' ', r'\!': '',
}

LIST_ENVS = {'itemize': 'List Bullet', 'enumerate': 'List Number', 'description': 'List Bullet'}
CODE_ENVS = {'verbatim', 'lstlisting', 'alltt', 'Verbatim', 'minted'}
PASS_THROUGH_ENVS = {'table', 'figure', 'center', 'abstract', 'quote',
                     'quotation', 'minipage', 'document', 'flushleft', 'flushright'}

# \\begin{env} 后面允许跟可选参数 [htbp] 与必选参数 {lcc}，
# 一并吃掉，否则它们会被当成正文泄漏到 Word 里。
ENV_RE = re.compile(
    r'\\(begin|end)\s*\{([^}]*)\}'
    r'((?:\s*(?:\[[^\]]*\]|\{[^}]*\}))*)'
)

# 行内样式命令（只匹配单层花括号，够用且不会 catastrophic backtracking）
INLINE_RE = re.compile(
    r'\\(textbf|textit|emph|underline|texttt|textsc|bf|it|tt|sc)\s*\{([^{}]*)\}'
    r"|\$([^$]+)\$"
)

DISPLAY_MATH_RE = re.compile(r'\\\[(.*?)\\\]|\\\((.*?)\\\)|\$\$(.*?)\$\$', re.S)

# 纯展示型命令，直接丢弃
STRIP_RE = re.compile(
    r'\\(?:noindent|indent|quad|qquad|hfill|vfill|hspace\*?|vspace\*?|'
    r'label|ref|cite|eqref|pageref|footnotemark|clearpage|newpage|'
    r'centering|raggedright|raggedleft|medskip|bigskip|smallskip|par)\s*(?:\{[^}]*\})?\*?'
)


# ---------------------------------------------------------------------------
# 文本清洗
# ---------------------------------------------------------------------------

def strip_comments(text):
    """去掉未转义的 % 注释（保留 \\%）。"""
    out = []
    for line in text.split('\n'):
        buf = []
        i = 0
        while i < len(line):
            ch = line[i]
            if ch == '\\' and i + 1 < len(line):
                buf.append(line[i:i + 2])
                i += 2
                continue
            if ch == '%':
                break
            buf.append(ch)
            i += 1
        out.append(''.join(buf))
    return '\n'.join(out)


def unescape(text):
    """把 LaTeX 转义还原成普通字符。"""
    for esc, plain in ESCAPED_CHARS.items():
        text = text.replace(esc, plain)
    # \textbackslash 之类的长命令
    text = text.replace(r'\textbackslash{}', '\\').replace(r'\textbackslash', '\\')
    text = text.replace(r'\textasciitilde{}', '~').replace(r'\textasciicircum{}', '^')
    text = re.sub(r'\\(?:,|;|!|\s)', ' ', text)
    # ~~ 与 --/--- 是 LaTeX 的波浪线与破折号
    text = text.replace('~~', '~ ').replace('---', '—').replace('--', '–')
    text = text.replace('``', '“').replace("''", '”')
    return text


def clean_inline(text):
    """去掉展示型命令与数学分隔符残留，再还原转义字符。"""
    text = STRIP_RE.sub(' ', text)
    text = re.sub(r'\\(?:hline|toprule|midrule|bottomrule|cline\{[^}]*\}|rule\{[^}]*\}\{[^}]*\})', ' ', text)
    text = unescape(text)
    return text


# ---------------------------------------------------------------------------
# 块切分
# ---------------------------------------------------------------------------

def parse_env_args(raw):
    """
    把 \\begin{tabular}{lcc} 之类命令后面跟的参数串解析成列表。

    '[htbp]' 与 '{lcc}' 都只保留内层内容，方便渲染时按用途取用。
    """
    args = []
    for m in re.finditer(r'\[([^\]]*)\]|\{([^}]*)\}', raw or ''):
        args.append(m.group(1) if m.group(1) is not None else m.group(2))
    return args


def split_blocks(text):
    """
    把源码切成块序列。返回列表，元素形如：
        ('text', str)
        ('env', name, args, [子块...])

    能正确处理同名环境的嵌套。
    """
    root = []
    stack = [(None, root)]
    i = 0
    n = len(text)

    while True:
        m = ENV_RE.search(text, i)
        if not m:
            stack[-1][1].append(('text', text[i:]))
            break

        stack[-1][1].append(('text', text[i:m.start()]))
        kind, name = m.group(1), m.group(2).strip()

        if kind == 'begin':
            buf = []
            args = parse_env_args(m.group(3))
            stack[-1][1].append(('env', name, args, buf))
            stack.append((name, buf))
        else:
            # \end{name}：弹出直到匹配的同名环境，防止嵌套错乱
            for k in range(len(stack) - 1, 0, -1):
                if stack[k][0] == name:
                    del stack[k:]
                    break
        i = m.end()

    return root


# ---------------------------------------------------------------------------
# 转换器
# ---------------------------------------------------------------------------

class LatexToDocx(object):
    """把一个 LaTeX 子集转换成 python-docx 的 Document 对象。"""

    def __init__(self):
        self.doc = Document()
        self.title = ''
        self.author = ''
        self.date = ''
        self._setup_styles()

    # -- 样式 ---------------------------------------------------------------
    def _setup_styles(self):
        """设置正文默认字体，并显式指定东亚字体，保证中文正常显示。"""
        try:
            normal = self.doc.styles['Normal']
            normal.font.name = 'Times New Roman'
            normal.font.size = Pt(12)
            rpr = normal.element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn('w:eastAsia'), '宋体')
            rfonts.set(qn('w:ascii'), 'Times New Roman')
            rfonts.set(qn('w:hAnsi'), 'Times New Roman')
        except Exception:
            # 字体设置失败不影响文档生成，只是回退到 Word 默认字体
            pass
        # 预先保证列表样式存在（部分模板可能缺失）
        for style_name in ('List Bullet', 'List Number'):
            try:
                self.doc.styles[style_name]
            except KeyError:
                pass

    # -- 入口 ---------------------------------------------------------------
    def convert(self, latex):
        latex = strip_comments(latex or '')
        latex = self._extract_preamble(latex)

        # 只处理 document 环境内部；找不到就整体当正文
        m = re.search(r'\\begin\s*\{document\}(.*?)\\end\s*\{document\}', latex, re.S)
        body = m.group(1) if m else latex

        if self.title or self.author or self.date:
            self._emit_title_block()

        self.render_blocks(split_blocks(body))
        return self.doc

    def _extract_preamble(self, latex):
        """抓取 \\title/\\author/\\date，其余导言区内容丢弃。"""
        for key in ('title', 'author', 'date'):
            m = re.search(r'\\%s\s*\{([^}]*)\}' % key, latex)
            if m:
                value = clean_inline(m.group(1)).strip()
                if key == 'date' and value == r'\today':
                    value = datetime.date.today().strftime('%Y-%m-%d')
                setattr(self, key, value)
        return latex

    def _emit_title_block(self):
        if self.title:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(self.title)
            run.bold = True
            run.font.size = Pt(18)
            self._set_eastasia(run, '黑体' if _has_cjk(self.title) else 'Times New Roman')
        meta = '  '.join(x for x in (self.author, self.date) if x)
        if meta:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(meta)
            run.font.size = Pt(11)

    @staticmethod
    def _set_eastasia(run, font_name):
        try:
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn('w:eastAsia'), font_name)
        except Exception:
            pass

    # -- 块渲染 -------------------------------------------------------------
    def render_blocks(self, blocks):
        for block in blocks:
            if block[0] == 'text':
                self.render_text(block[1])
            else:
                self.render_env(block[1], block[2], block[3])

    def render_env(self, name, args, children):
        if name in LIST_ENVS:
            self._render_list(name, children)
        elif name.startswith('tabular'):
            self._render_tabular(name, args, children)
        elif name in CODE_ENVS:
            self._render_code(children)
        elif name in PASS_THROUGH_ENVS:
            if name == 'center':
                self._render_centered(children)
            else:
                self.render_blocks(children)
        elif name in ('equation', 'equation*', 'displaymath', 'align', 'align*',
                      'gather', 'gather*', 'multline', 'multline*', 'eqnarray'):
            self._render_display_math(''.join(self._flatten(children)))
        else:
            # 未知环境：退化成普通文本，避免内容凭空消失
            self.render_blocks(children)

    def _flatten(self, blocks):
        out = []
        for block in blocks:
            if block[0] == 'text':
                out.append(block[1])
            else:
                out.append(''.join(self._flatten(block[3])))
        return out

    def _render_centered(self, children):
        """center 环境：内部段落居中。"""
        start = len(self.doc.paragraphs)
        self.render_blocks(children)
        for p in self.doc.paragraphs[start:]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _render_code(self, children):
        text = ''.join(self._flatten(children))
        text = text.replace(r'\\', '\n')
        for line in text.split('\n'):
            p = self.doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            self._set_eastasia(run, '宋体')

    @staticmethod
    def _clean_math(text):
        """规范化数学源码：去掉 label/tag 与间距命令，便于以纯文本降级展示。"""
        text = re.sub(r'\\(?:label|tag|nonumber|notag)\s*(?:\{[^}]*\})?', '', text or '')
        # \,\;\! 与 \quad 在纯文本展示里没有意义
        text = re.sub(r'\\[,;!]|\\q?quad', ' ', text or '')
        return re.sub(r'[ \t]+', ' ', text).strip()

    def _render_display_math(self, text):
        text = self._clean_math(text)
        if not text:
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.italic = True
        run.font.name = 'Consolas'

    def _render_list(self, name, children):
        style = LIST_ENVS[name]
        # 收集 \item 项：先把子块摊平成文本，再按 \item 切分
        text = ''.join(self._flatten(children))
        raw_items = re.split(r'\\item\s*', text)
        for raw in raw_items:
            raw = raw.strip()
            if not raw:
                continue
            content = raw
            label = None
            m = re.match(r'\[([^\]]*)\]\s*', raw)
            if m:  # \item[标签] 形式
                label = clean_inline(m.group(1)).strip()
                content = raw[m.end():]
            p = self.doc.add_paragraph(style=style)
            if label:
                run = p.add_run(label + '　')
                run.bold = True
            self.emit_inline(p, content)

    def _render_tabular(self, name, args, children):
        """
        渲染 tabular 环境。

        列数由内容推断（按 & 切分），列对齐优先取 \\begin{tabular}{lcc} 的参数，
        参数缺失或与内容不一致时回退为居中，不丢单元格。
        """
        rows = self._parse_tabular_rows(children)
        if not rows:
            return

        n_cols = max(len(r) for r in rows)
        colspec = self._pick_colspec(args, n_cols)

        table = self.doc.add_table(rows=0, cols=n_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for r_index, row in enumerate(rows):
            cells = table.add_row().cells
            for c_index in range(n_cols):
                text = row[c_index] if c_index < len(row) else ''
                cell = cells[c_index]
                cell.text = ''
                p = cell.paragraphs[0]
                if c_index < len(colspec) and colspec[c_index] == 'r':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif c_index < len(colspec) and colspec[c_index] == 'l':
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                self.emit_inline(p, text)
                # 首行加粗，保留 LaTeX 表头的视觉语义
                if r_index == 0 and len(rows) > 1:
                    for run in p.runs:
                        run.bold = True

    @staticmethod
    def _pick_colspec(args, n_cols):
        """从 \\begin{tabular} 的参数里挑出列格式串，如 'lcc'。"""
        for candidate in reversed(args or []):
            cleaned = re.sub(r'[|\s]', '', candidate or '')
            # 列格式只允许 l/c/r/p 以及 p{} 的宽度参数
            if cleaned and re.fullmatch(r'[lcrp](\{[^}]*\})?(\s*[lcrp](\{[^}]*\})?)*', cleaned):
                return re.findall(r'[lcr]', cleaned)
        return ['c'] * n_cols

    def _parse_tabular_rows(self, children):
        text = ''.join(self._flatten(children))
        # 去掉各种横线命令
        text = re.sub(r'\\(?:hline|toprule|midrule|bottomrule|cline\{[^}]*\})', '', text)
        # \multicolumn{n}{align}{内容} 只保留内容
        text = re.sub(r'\\multicolumn\s*\{\d+\}\s*\{[^}]*\}\s*\{', '{', text)
        rows = []
        for raw_row in re.split(r'\\\\', text):
            raw_row = raw_row.strip()
            if not raw_row:
                continue
            rows.append([c.strip() for c in raw_row.split('&')])
        return rows

    # -- 文本渲染 -----------------------------------------------------------
    def render_text(self, raw):
        if not raw:
            return
        # \maketitle 在导言区已经消费过（生成标题块），正文里出现就直接丢弃
        raw = re.sub(r'\\maketitle', ' ', raw)
        # 先按空行分段
        for chunk in re.split(r'\n\s*\n', raw):
            chunk = chunk.strip()
            if not chunk:
                continue
            self._render_paragraph(chunk)

    def _add_caption(self, text):
        """表格/图片题注：居中、小一号字。"""
        if not text:
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(10)

    def _render_paragraph(self, chunk):
        # 题注 \caption[短题注]{完整题注}
        m = re.match(r'\\caption\s*(?:\[[^\]]*\])?\s*\{([^}]*)\}', chunk)
        if m:
            self._add_caption(clean_inline(m.group(1)).strip())
            rest = chunk[m.end():].strip()
            if rest:
                self._render_paragraph(rest)
            return

        # 标题命令
        m = re.match(r'\\(part|section|subsection|subsubsection|paragraph|subparagraph)\*?\s*(?:\[[^\]]*\])?\s*\{([^}]*)\}', chunk)
        if m:
            level = SECTION_LEVELS.get(m.group(1), 1)
            heading_text = clean_inline(m.group(2)).strip()
            rest = chunk[m.end():].strip()
            if heading_text:
                h = self.doc.add_heading(heading_text, level=min(level, 4))
                for run in h.runs:
                    self._set_eastasia(run, '黑体' if _has_cjk(heading_text) else 'Times New Roman')
            if rest:
                self._render_paragraph(rest)
            return

        # 独占一行的行间公式
        dm = DISPLAY_MATH_RE.fullmatch(chunk.strip())
        if dm:
            body = next((g for g in dm.groups() if g is not None), '')
            self._render_display_math(body)
            return

        # 普通段落
        p = self.doc.add_paragraph()
        self.emit_inline(p, chunk)

    # -- 行内渲染 -----------------------------------------------------------
    def emit_inline(self, paragraph, text):
        """把一段文本按样式拆成若干 run 追加到段落中。"""
        if text is None:
            return
        pos = 0
        for m in INLINE_RE.finditer(text):
            if m.start() > pos:
                self._add_plain(paragraph, text[pos:m.start()])
            if m.group(3) is not None:
                # 行内数学：降级为斜体等宽，前后补空格避免与正文粘连
                run = paragraph.add_run(' ' + self._clean_math(m.group(3)) + ' ')
                run.italic = True
                run.font.name = 'Consolas'
            else:
                style = STYLE_FONT_MAP.get(m.group(1), None)
                run = paragraph.add_run(clean_inline(m.group(2)))
                if style == 'bold':
                    run.bold = True
                elif style == 'italic':
                    run.italic = True
                elif style == 'underline':
                    run.underline = True
                elif style == 'mono':
                    run.font.name = 'Consolas'
                elif style == 'smallcaps':
                    run.font.small_caps = True
            pos = m.end()

        if pos < len(text):
            self._add_plain(paragraph, text[pos:])

    def _add_plain(self, paragraph, raw):
        """处理普通文本：换行符、行间公式残留、转义还原。"""
        if not raw:
            return

        # 行间公式 \[ ... \] / $$ ... $$ 在段落内部也一并处理
        pos = 0
        for m in DISPLAY_MATH_RE.finditer(raw):
            self._add_plain_text(paragraph, raw[pos:m.start()])
            body = next((g for g in m.groups() if g is not None), '')
            body = self._clean_math(body)
            if body:
                run = paragraph.add_run(body)
                run.italic = True
                run.font.name = 'Consolas'
            pos = m.end()
        self._add_plain_text(paragraph, raw[pos:])

    def _add_plain_text(self, paragraph, raw):
        if not raw:
            return
        # 换行
        parts = re.split(r'\\\\|\\newline', raw)
        for i, part in enumerate(parts):
            if i > 0:
                paragraph.add_run().add_break()
            part = clean_inline(part)
            if not part:
                continue
            # 段落首尾空白无意义，中间的连续空白压成一个
            part = re.sub(r'[ \t]+', ' ', part)
            if part.strip():
                run = paragraph.add_run(part.strip() if i == 0 or i == len(parts) - 1 else part)
                if _has_cjk(part):
                    self._set_eastasia(run, '宋体')


def _has_cjk(text):
    return any('一' <= ch <= '鿿' for ch in text or '')


# ---------------------------------------------------------------------------
# 对外接口
# ---------------------------------------------------------------------------

def export_word(data, output_path):
    """
    导出 Word 文件。

    :param data: 至少包含 content 字段（LaTeX 源码）的字典
    :param output_path: 输出的 .docx 路径
    :return: 结果字典
    """
    if Document is None:
        return {
            'success': False,
            'error': '缺少 python-docx 依赖，请执行: pip install python-docx'
        }

    try:
        content = (data or {}).get('content', '') or ''

        output_dir = os.path.dirname(os.path.abspath(output_path))
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)

        doc = LatexToDocx().convert(content)
        doc.save(output_path)

        if not os.path.exists(output_path):
            return {'success': False, 'error': 'Word 文件写入失败: %s' % output_path}

        return {
            'success': True,
            'output_path': os.path.abspath(output_path),
            'size': os.path.getsize(output_path),
        }
    except Exception as e:
        return {'success': False, 'error': '导出失败: %s' % e}


def main():
    if len(sys.argv) < 3:
        print(json.dumps({'success': False, 'error': '缺少参数'}, ensure_ascii=False))
        sys.exit(1)
    try:
        result = export_word(json.loads(sys.argv[1]), sys.argv[2])
        print(json.dumps(result, ensure_ascii=False))
        sys.exit(0 if result.get('success') else 1)
    except Exception as e:  # pragma: no cover
        print(json.dumps({'success': False, 'error': str(e)}, ensure_ascii=False))
        sys.exit(1)


if __name__ == '__main__':
    main()

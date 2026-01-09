"""
Word Exporter
Exports documents to Word format using python-docx.
"""

from docx import Document
from docx.shared import Inches, Pt


class WordExporter:
    """Handles Word export functionality."""

    def __init__(self):
        self.document = None

    def export(self, data, output_path):
        """
        Export data to Word document.

        Args:
            data: Document data structure
            output_path: Output Word file path

        Returns:
            tuple: (success: bool, error: str)
        """
        try:
            self.document = Document()
            content = data.get('content', '')

            if content:
                self._parse_and_add_content(content)

            self.document.save(output_path)
            return (True, '')
        except Exception as e:
            return (False, str(e))

    def _parse_and_add_content(self, content):
        """Parse LaTeX-like content and add to Word document."""
        lines = content.split('\n')
        in_table = False

        for line in lines:
            line = line.strip()

            if r'\begin{tabular}' in line:
                in_table = True
                table_data = self._extract_table_data(lines, lines.index(line))
                self._add_table(table_data)
            elif r'\end{tabular}' in line:
                in_table = False
            elif not in_table and line and not line.startswith('\\'):
                self.document.add_paragraph(line)

    def _extract_table_data(self, lines, start_index):
        """Extract table data from LaTeX content."""
        table_data = []
        current_row = []

        for line in lines[start_index:]:
            if r'\end{tabular}' in line:
                if current_row:
                    table_data.append(current_row)
                break
            elif r'\\' in line and not line.startswith('\\'):
                if current_row:
                    table_data.append(current_row)
                cells = [cell.strip() for cell in line.split('&')]
                current_row = cells
            elif line.strip():
                cells = [cell.strip() for cell in line.split('&')]
                current_row = cells

        return table_data

    def _add_table(self, table_data):
        """Add table to Word document."""
        if not table_data:
            return

        table = self.document.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Table Grid'

        for i, row_data in enumerate(table_data):
            for j, cell_data in enumerate(row_data):
                cell = table.rows[i].cells[j]
                cell.text = cell_data

    def _add_formula(self, formula_data):
        """Add formula to Word document."""
        formula_text = formula_data.get('text', '')
        if formula_text:
            self.document.add_paragraph(formula_text)


if __name__ == '__main__':
    if len(sys.argv) > 2:
        try:
            data = json.loads(sys.argv[1])
            output_path = sys.argv[2]

            exporter = WordExporter()
            success, error = exporter.export(data, output_path)

            if success:
                print(json.dumps({
                    'success': True,
                    'message': 'Word document exported successfully'
                }))
            else:
                print(json.dumps({
                    'success': False,
                    'error': error
                }))

        except Exception as e:
            print(json.dumps({
                'success': False,
                'error': str(e)
            }))